from __future__ import annotations

import io
import os
import re
from zipfile import ZipFile

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


SOURCE = "/app/output/source_TMHCC_AI_PI_EO_Board_Paper_Final_Board_Ready.docx"
OUT_PDF = "/app/output/TMHCC_AI_PI_EO_Board_Paper_Rebuilt_13_May_2026.pdf"
OUT_DOCX = "/app/output/TMHCC_AI_PI_EO_Board_Paper_Rebuilt_13_May_2026.docx"

NAVY = colors.HexColor("#071B33")
GOLD = colors.HexColor("#B68A35")
SLATE = colors.HexColor("#3D4656")
PALE = colors.HexColor("#F4F1EA")
LINE = colors.HexColor("#C9CED6")


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)


def clean(text: str) -> str:
    text = text.replace("&amp;", "&")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def esc(text: str) -> str:
    return (
        clean(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def paragraph_image_paths(paragraph: DocxParagraph, document: Document, media_dir: str) -> list[str]:
    embeds = paragraph._p.xpath(".//a:blip/@r:embed")
    paths = []
    for rid in embeds:
        part = document.part.related_parts[rid]
        name = os.path.basename(part.partname)
        out = os.path.join(media_dir, name)
        if not os.path.exists(out):
            with open(out, "wb") as f:
                f.write(part.blob)
        paths.append(out)
    return paths


def build_styles():
    base = getSampleStyleSheet()
    base.add(ParagraphStyle(
        name="CoverTitle", fontName="Helvetica-Bold", fontSize=28, leading=32,
        textColor=NAVY, alignment=TA_LEFT, spaceAfter=20,
    ))
    base.add(ParagraphStyle(
        name="CoverSub", fontName="Helvetica", fontSize=12.5, leading=17,
        textColor=SLATE, alignment=TA_LEFT, spaceAfter=8,
    ))
    base.add(ParagraphStyle(
        name="H1Board", fontName="Helvetica-Bold", fontSize=18, leading=22,
        textColor=NAVY, spaceBefore=18, spaceAfter=9, keepWithNext=True,
    ))
    base.add(ParagraphStyle(
        name="H2Board", fontName="Helvetica-Bold", fontSize=13.5, leading=17,
        textColor=SLATE, spaceBefore=12, spaceAfter=6, keepWithNext=True,
    ))
    base.add(ParagraphStyle(
        name="BodyBoard", fontName="Helvetica", fontSize=9.3, leading=13.0,
        textColor=colors.HexColor("#27313F"), spaceAfter=6,
    ))
    base.add(ParagraphStyle(
        name="SmallBoard", fontName="Helvetica", fontSize=8.0, leading=10.5,
        textColor=SLATE, spaceAfter=4,
    ))
    base.add(ParagraphStyle(
        name="TOCBoard", fontName="Helvetica", fontSize=10.5, leading=15,
        textColor=colors.HexColor("#27313F"), spaceAfter=3,
    ))
    base.add(ParagraphStyle(
        name="TableCell", fontName="Helvetica", fontSize=7.4, leading=9.4,
        textColor=colors.HexColor("#27313F"), spaceAfter=0,
    ))
    base.add(ParagraphStyle(
        name="TableHead", fontName="Helvetica-Bold", fontSize=7.6, leading=9.6,
        textColor=colors.white, spaceAfter=0,
    ))
    return base


def draw_page(canvas, doc):
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(LINE)
    canvas.setLineWidth(0.6)
    canvas.line(doc.leftMargin, height - 1.35 * cm, width - doc.rightMargin, height - 1.35 * cm)
    canvas.setFillColor(NAVY)
    canvas.setFont("Helvetica-Bold", 8)
    canvas.drawString(doc.leftMargin, height - 1.08 * cm, "TMHCC | AI PI/E&O Board Paper | UK-first AI endorsement strategy")
    canvas.setFillColor(SLATE)
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(width - doc.rightMargin, 0.85 * cm, f"TMHCC Confidential | Page {doc.page}")
    canvas.restoreState()


def docx_table_to_reportlab(table: DocxTable, styles) -> Table:
    rows = []
    for r, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            txt = "\n".join(clean(p.text) for p in cell.paragraphs if clean(p.text))
            style = styles["TableHead"] if r == 0 else styles["TableCell"]
            cells.append(Paragraph(esc(txt), style))
        rows.append(cells)
    if not rows:
        return Spacer(1, 1)
    ncols = max(len(r) for r in rows)
    # Repeat malformed cells defensively so ReportLab receives a rectangle.
    for row in rows:
        while len(row) < ncols:
            row.append(Paragraph("", styles["TableCell"]))
    widths = [16.8 * cm / ncols] * ncols
    t = Table(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, PALE]),
    ]))
    return t


def build_pdf():
    document = Document(SOURCE)
    os.makedirs("/app/output/pdf_media", exist_ok=True)

    styles = build_styles()
    doc = SimpleDocTemplate(
        OUT_PDF, pagesize=A4,
        leftMargin=1.85 * cm, rightMargin=1.85 * cm,
        topMargin=1.75 * cm, bottomMargin=1.45 * cm,
        title="TMHCC AI PI/E&O Board Paper",
        author="TMHCC",
    )
    story = []

    # Cover page
    story.append(Spacer(1, 2.2 * cm))
    story.append(Paragraph("TMHCC", ParagraphStyle("Brand", fontName="Helvetica-Bold", fontSize=18, textColor=GOLD, spaceAfter=18)))
    story.append(Paragraph("AI in PI, E&O and Media Liability", styles["CoverTitle"]))
    story.append(Paragraph("Board-ready final paper: UK-first exclusion-plus-write-back strategy for IT and technology, construction and built environment, and media, music, events, film and television", styles["CoverSub"]))
    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("Date: 13 May 2026", styles["CoverSub"]))
    story.append(Paragraph("Prepared for: TMHCC Board review", styles["CoverSub"]))
    story.append(Paragraph("Document status: Board-ready working paper for legal, underwriting and product governance sign-off", styles["CoverSub"]))
    story.append(Spacer(1, 1.0 * cm))
    story.append(Table([[Paragraph("Recommended decision", styles["TableHead"])], [Paragraph("Approve a UK-first AI endorsement strategy that starts with IT and technology plus selected media/creative accounts, then extends to construction and built-environment professionals once the governance questionnaire, rating model and claims coding are operational.", styles["BodyBoard"])]], colWidths=[16.8 * cm], style=TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("BACKGROUND", (0, 1), (-1, -1), PALE),
        ("BOX", (0, 0), (-1, -1), 0.7, GOLD),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ])))
    story.append(PageBreak())

    # PDF TOC generated from headings.
    headings = []
    for p in document.paragraphs:
        if p.style.name in ("Heading 1", "Heading 2") and clean(p.text):
            level = 1 if p.style.name == "Heading 1" else 2
            headings.append((level, clean(p.text)))
    story.append(Paragraph("Contents", styles["H1Board"]))
    for level, title in headings:
        indent = "&nbsp;&nbsp;&nbsp;&nbsp;" if level == 2 else ""
        story.append(Paragraph(f"{indent}{esc(title)}", styles["TOCBoard"]))
    story.append(PageBreak())

    # Main body starts at first Heading 1.
    started = False
    for block in iter_block_items(document):
        if isinstance(block, DocxParagraph):
            text = clean(block.text)
            if block.style.name == "Heading 1" and text:
                started = True
            if not started:
                continue

            imgs = paragraph_image_paths(block, document, "/app/output/pdf_media")
            if imgs:
                for img_path in imgs:
                    img = Image(img_path)
                    img._restrictSize(16.4 * cm, 10.6 * cm)
                    story.append(KeepTogether([img, Spacer(1, 0.15 * cm)]))
                continue

            if not text:
                story.append(Spacer(1, 0.1 * cm))
                continue

            if block.style.name == "Heading 1":
                story.append(Paragraph(esc(text), styles["H1Board"]))
            elif block.style.name == "Heading 2":
                story.append(Paragraph(esc(text), styles["H2Board"]))
            else:
                if text.startswith("Source:") or text.startswith("Drafting note:") or text.startswith("Plain-English"):
                    story.append(Paragraph(esc(text), styles["SmallBoard"]))
                elif text.startswith("ENDORSEMENT") or re.match(r"^[A-E]\. ", text):
                    story.append(Paragraph(esc(text), styles["H2Board"]))
                else:
                    story.append(Paragraph(esc(text), styles["BodyBoard"]))
        elif isinstance(block, DocxTable) and started:
            story.append(docx_table_to_reportlab(block, styles))
            story.append(Spacer(1, 0.18 * cm))

    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)


if __name__ == "__main__":
    # Keep the Word document intact, as it already contains the automated TOC, headers, footers and embedded charts.
    with open(SOURCE, "rb") as src, open(OUT_DOCX, "wb") as dst:
        dst.write(src.read())
    build_pdf()
    print(OUT_DOCX)
    print(OUT_PDF)